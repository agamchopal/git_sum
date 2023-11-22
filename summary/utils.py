import nltk
nltk.download('stopwords')
nltk.download('punkt')
nltk.data.path.append("path_to_nltk_data")
from nltk.corpus import stopwords

stopwords.words('english')


from django.shortcuts import render
import spacy
from nltk.corpus import stopwords


nlp = spacy.load('en_core_web_sm')
import nltk
nltk.download('stopwords')
nltk.download('punkt')
from nltk.tokenize import word_tokenize, sent_tokenize
stopWords = set(stopwords.words("english"))

def search(pat, txt, q):
        d = 256
        count = 0
        M = len(pat)
        N = len(txt)
        i = 0
        j = 0
        p = 0  
        t = 0  
        h = 1
       
        for i in range(M - 1):
            h = (h * d) % q
       
        for i in range(M):
            p = (d * p + ord(pat[i])) % q
            t = (d * t + ord(txt[i])) % q
       
        for i in range(N - M + 1):
        
            if p == t:
                
                for j in range(M):
                    if txt[i + j] != pat[j]:
                        break
                    else:
                        j += 1
                
                if j == M:
                    count += 1
          
            if i < N - M:
                t = (d * (t - ord(txt[i]) * h) + ord(txt[i + M])) % q
                
                if t < 0:
                    t = t + q
        return count

def robin(keywords,text, summary_length):
    q = 101
    newdict = {}
    keywords = list(map(str, keywords.split(',')))
    for i in keywords:
        newdict[i] = 0

    for pat in keywords:
        try:
            newdict[pat] = search(pat.lower(), text.lower(), q)
        except:
            pass

    doc = nlp(text)
    text = []
    for i in doc.sents:
        i = str(i)
        message = i.split('\n.')
        text.append(message[0])
    relevant_sentences = []

   
    for sentence in text:
        sentence = str(sentence)
        added = False
        for keyword in keywords:
            if keyword.lower() in sentence.lower() and not added:
                relevant_sentences.append(sentence.strip())
                added = True  

    '''for sentence in text:
        sentence = str(sentence)
        for keyword in keywords:
            if keyword.lower() in sentence.lower():
        #if l.lower() in sentence.lower():
                relevant_sentences.append(sentence.strip())'''

    document1 = "".join(relevant_sentences)

    words = word_tokenize(document1)
    
    freqTable = dict()
    for word in words:
        word = word.lower()
        if word in stopWords:
            continue
        if word in freqTable:
            freqTable[word] += 1
        else:
            freqTable[word] = 1
   
    sentences = sent_tokenize(document1)
    sentenceValue = dict()
    for sentence in sentences:
        for word, freq in freqTable.items():
            if word in sentence.lower():
                if sentence in sentenceValue:
                    sentenceValue[sentence] += freq
                else:
                    sentenceValue[sentence] = freq

    sumValues = 0
    for sentence in sentenceValue:
        sumValues += sentenceValue[sentence]
    
    average = 1
    if len(sentenceValue) != 0:
        average = int(sumValues / len(sentenceValue))
   
    summary = ''
    if summary_length == 'one_third':
        target_length = len(sentences) // 3
    elif summary_length == 'two_thirds':
        target_length = 2 * (len(sentences) // 3)
    else:
        target_length = len(sentences)
    
    current_length = 0
    for sentence in sentences:
        if sentence in sentenceValue and current_length < target_length:
            #summary += " " + sentence
            #current_length += 1
            for keyword in keywords:
                #sentence = sentence.replace(keyword,f'<span style="font-weight: bold;" class="text-danger">{keyword}</span>')
                sentence = sentence.replace(keyword, f'<span style="font-weight: bold; background-color: yellow;">{keyword}</span>')

            summary += " " +sentence
            current_length += 1
            
            

            




    


    '''
    for sentence in sentences: 
        if (sentence in sentenceValue) and (sentenceValue[sentence] > (1.2 * average)):
            summary += " " + sentence
    
    '''
    summary = re.sub(r'(?<=[a-zA-Z0-9])\.(?=[a-zA-Z])', '. ', summary)
    #print(summary)
    
    
    

    

    lst = set()
    for i in newdict.items():
        lst.add(i[0])

    """summary Statistic part"""
    
    summary_count = count_paragraphs_helper(summary)
    summary_lines = count_lines_helper(summary)
    summary_updated_lines = count_lines(summary)
    summary_char_count = character_count(summary)
    summary_char_count_one = count_char_one(summary)
    summary_word_count = count_words(summary)
    summary_pages = count_pages(summary)


    params = {   'sum': sorted(newdict.items(), key=lambda kv:(kv[1], kv[0])), 
              
                 'Summary': summary, 
                 'Text':text[0], 
                 "Keywords": keywords,

                   

                'summary_count': summary_count,
                'summary_lines': summary_lines,
                
                'summary_char_count': summary_char_count,
                'summary_char_count_one': summary_char_count_one,
                'summary_word_count': summary_word_count,
                'summary_pages': summary_pages
                 
                 
                 
                 
                 }
    
    
    return params
from docx import Document
from io import BytesIO
from django.http import HttpResponse

def download_summary_as_docx(summary,filename):
    # Create a new Document object
    doc = Document()

    # Add a title to the document
    doc.add_heading('Summary', level=1)

    # Add the summary to the document
    doc.add_paragraph(summary)

    # Create an in-memory output stream for the document
    f = BytesIO()

    # Save the document to the in-memory stream
    doc.save(f)
    length = f.tell()
    f.seek(0)

    # Create an HttpResponse with the content type set to 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    # Set the Content-Disposition header to indicate the filename
    response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
    
    # Set the Content-Length header
    response['Content-Length'] = length

    return response




from django.http import HttpResponse
from docx import Document
import io

'''def download_summary(request):
    # Assuming you have the keywords and text available in your request
    keywords = request.POST.get('keywords', '')  # Adjust the key based on how you send the data
    text = request.POST.get('text', '')  # Adjust the key based on how you send the data
    summary_length = request.POST.get('summary_length', '')  # Adjust the key based on how you send the data
    
    # Call your robin function to get the params
    params = robin(keywords, text, summary_length)
    
    # Get the summary from the params
    summary = params['Summary']
    
    # Generate DOCX file
    doc = Document()
    doc.add_paragraph(summary)
    
    # In-memory output stream for the document
    f = io.BytesIO()
    doc.save(f)
    length = f.tell()
    f.seek(0)

    response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="auto_generated_summary.doc"'
    response['Content-Length'] = length
    return response

'''
    


"""summary Statistic part"""
import re

def count_paragraphs_helper(text):
    paragraphs = re.split(r'\n\s*\n', text.strip())
    return len(paragraphs)

def count_lines_helper(text):
    sentences = re.split(r'[.!?]\s*|(?<!\w)[,;](?!\w)\s*|$' , text)
    non_empty_sentences = [sentence for sentence in sentences if sentence.strip() != '']
    return len(non_empty_sentences)


    
    line_count = len(sentences)
    return line_count

def count_pages(text):
    
    lines_per_page = 50
    lines = count_lines_helper(text)
    pages = lines // lines_per_page
    if lines % lines_per_page != 0:
        pages += 1
    return pages

def character_count(text):
    c = 0
    for i in range(len(text)):
        if text[i] != " ":
            c += 1
    return c

def count_char_one(text):
    return len(text)

def count_words(text):
    words = text.split()
    return len(words)

def count_lines(text):
    lines = re.split(r'\.\s+', text)
    formatted_lines = '<br>'.join(lines)
    return formatted_lines
from docx import Document
from django.http import HttpResponse
import io
def robinn(keywords,text, summary_length):
    q = 101
    newdict = {}
    keywords = list(map(str, keywords.split(',')))
    for i in keywords:
        newdict[i] = 0

    for pat in keywords:
        try:
            newdict[pat] = search(pat.lower(), text.lower(), q)
        except:
            pass

    doc = nlp(text)
    text = []
    for i in doc.sents:
        i = str(i)
        message = i.split('\n.')
        text.append(message[0])
    relevant_sentences = []

   
    for sentence in text:
        sentence = str(sentence)
        added = False
        for keyword in keywords:
            if keyword.lower() in sentence.lower() and not added:
                relevant_sentences.append(sentence.strip())
                added = True  

    '''for sentence in text:
        sentence = str(sentence)
        for keyword in keywords:
            if keyword.lower() in sentence.lower():
        #if l.lower() in sentence.lower():
                relevant_sentences.append(sentence.strip())'''

    document1 = "".join(relevant_sentences)

    words = word_tokenize(document1)
    
    freqTable = dict()
    for word in words:
        word = word.lower()
        if word in stopWords:
            continue
        if word in freqTable:
            freqTable[word] += 1
        else:
            freqTable[word] = 1
   
    sentences = sent_tokenize(document1)
    sentenceValue = dict()
    for sentence in sentences:
        for word, freq in freqTable.items():
            if word in sentence.lower():
                if sentence in sentenceValue:
                    sentenceValue[sentence] += freq
                else:
                    sentenceValue[sentence] = freq

    sumValues = 0
    for sentence in sentenceValue:
        sumValues += sentenceValue[sentence]
    
    average = 1
    if len(sentenceValue) != 0:
        average = int(sumValues / len(sentenceValue))
   
    summary = ''
    if summary_length == 'one_third':
        target_length = len(sentences) // 3
    elif summary_length == 'two_thirds':
        target_length = 2 * (len(sentences) // 3)
    else:
        target_length = len(sentences)
    
    current_length = 0
    for sentence in sentences:
        if sentence in sentenceValue and current_length < target_length:
            #summary += " " + sentence
            #current_length += 1
            for keyword in keywords:
                sentence = sentence.replace(keyword,f'<span style="font-weight: bold;" class="text-danger">{keyword}</span>')
            summary += " " 
            current_length += 1
            
            

            




    


    '''
    for sentence in sentences: 
        if (sentence in sentenceValue) and (sentenceValue[sentence] > (1.2 * average)):
            summary += " " + sentence
    
    '''
    summary = re.sub(r'(?<=[a-zA-Z0-9])\.(?=[a-zA-Z])', '. ', summary)
    #print(summary)
    
    
    

    

    lst = set()
    for i in newdict.items():
        lst.add(i[0])

    """summary Statistic part"""
    
    summary_count = count_paragraphs_helper(summary)
    summary_lines = count_lines_helper(summary)
    summary_updated_lines = count_lines(summary)
    summary_char_count = character_count(summary)
    summary_char_count_one = count_char_one(summary)
    summary_word_count = count_words(summary)
    summary_pages = count_pages(summary)


    params = {   'sum': sorted(newdict.items(), key=lambda kv:(kv[1], kv[0])), 
              
                 'Summary': summary, 
                 'Text':text[0], 
                 "Keywords": keywords,

                   

                'summary_count': summary_count,
                'summary_lines': summary_lines,
                
                'summary_char_count': summary_char_count,
                'summary_char_count_one': summary_char_count_one,
                'summary_word_count': summary_word_count,
                'summary_pages': summary_pages
                 
                 
                 
                 
                 }
    
    
    return params




